"""
File Name: test_resume_extraction.py
Author: Senthilnathan Karuppaiah
Date: 11-JUL-2026
Description:
Resume-extraction eval (no LLM) — format allowlist, size cap, scanned
rejection, and normalized output against the committed synthetic fixture.

This suite pins extraction by:
1. The synthetic PDF fixture extracts its fictional identity, wrapped
   at 120 columns.
2. .doc and unknown extensions are rejected with clear errors.
3. Near-empty content trips the min-words (no-OCR) guard; oversized
   files trip the byte cap before reading.
"""

# Import necessary libraries
from pathlib import Path

import pytest

from job_matcher.resume import ResumeError, extract_resume_text

FIXTURES = Path(__file__).parent.parent / "evals" / "data"


def test_synthetic_pdf_extracts_identity():
    text = extract_resume_text(FIXTURES / "resume" / "synthetic-resume.pdf")
    assert "Jordan Rivera".lower() in text.lower()
    assert "jordan.rivera@example.com" in text.lower()
    assert all(len(line) <= 120 for line in text.splitlines())


def test_txt_and_md_pass_through(tmp_path):
    p = tmp_path / "resume.md"
    p.write_text("# Jane Doe\n" + ("data engineer with python and sql experience " * 10))
    assert "Jane Doe" in extract_resume_text(p)


def test_docx_extracts(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Casey Kim — Platform Engineer")
    doc.add_paragraph("Built streaming pipelines with Kafka and Spark. " * 12)
    p = tmp_path / "resume.docx"
    doc.save(str(p))
    assert "Casey Kim" in extract_resume_text(p)


def test_legacy_doc_rejected(tmp_path):
    p = tmp_path / "resume.doc"
    p.write_bytes(b"old word binary")
    with pytest.raises(ResumeError, match=r"\.doc"):
        extract_resume_text(p)


def test_unknown_extension_rejected(tmp_path):
    p = tmp_path / "resume.rtf"
    p.write_text("hello")
    with pytest.raises(ResumeError, match="unsupported"):
        extract_resume_text(p)


def test_min_words_guard_rejects_near_empty(tmp_path):
    p = tmp_path / "resume.txt"
    p.write_text("just five words right here")
    with pytest.raises(ResumeError, match="no extractable text"):
        extract_resume_text(p)


def test_byte_cap(tmp_path, monkeypatch):
    monkeypatch.setenv("MAX_RESUME_BYTES", "100")
    p = tmp_path / "resume.txt"
    p.write_text("word " * 200)
    with pytest.raises(ResumeError, match="byte cap"):
        extract_resume_text(p)


def test_missing_file_rejected(tmp_path):
    with pytest.raises(ResumeError, match="not found"):
        extract_resume_text(tmp_path / "nope.pdf")
