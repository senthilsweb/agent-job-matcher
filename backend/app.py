"""
Job Fit Analyzer - Production Version
Clean, consolidated app with deterministic scoring and improved UI
Usage: streamlit run app.py
"""
import streamlit as st
import requests
from bs4 import BeautifulSoup
from pathlib import Path
from typing import List, Optional, Dict, Any
from io import BytesIO
from enum import Enum
from datetime import datetime
from pydantic import BaseModel, Field
from pydantic_ai import Agent
from pypdf import PdfReader
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from fpdf import FPDF
import os
import time
import json
import re
from urllib.parse import urlparse


# ============================================================================
# CONFIGURATION
# ============================================================================

BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
PROMPTS_DIR = BASE_DIR / "prompts"

CANDIDATE_INFO = {
    "name": "SENTHILNATHAN KARUPPAIAH",
    "title": "Data Governance & Privacy · Forward-Deployed Engineer · Data & GenAI Architect",
    "email": "nathansweb@icloud.com",
    "phone": "732-955-1143",
    "location": "New Jersey, USA",
    "linkedin": "linkedin.com/in/senthilsweb",
    "github": "github.com/senthilsweb",
    "website": "senthilsweb.com"
}

# Deterministic Scoring Weights
SCORE_WEIGHTS = {
    "required_skills_match": 40,      # % of required skills matched
    "preferred_skills_match": 20,     # % of preferred skills matched
    "experience_years_match": 20,     # Experience level alignment
    "domain_expertise_match": 20,     # Domain/industry alignment
}


# ============================================================================
# TEMPLATE ENGINE (Mustache-style)
# ============================================================================

def render_template(template: str, data: Dict[str, Any]) -> str:
    """Render mustache-style template with {{variable}} placeholders."""
    def replace(match):
        key = match.group(1).strip()
        return str(data.get(key, f"{{{{{{key}}}}}}"))
    return re.sub(r'\{\{(\s*\w+\s*)\}\}', replace, template)


def load_template(name: str) -> str:
    """Load template from templates folder."""
    path = TEMPLATES_DIR / name
    if path.exists():
        return path.read_text(encoding="utf-8")
    raise FileNotFoundError(f"Template not found: {name}")


def load_prompt(name: str) -> str:
    """Load prompt from prompts folder."""
    path = PROMPTS_DIR / name
    if path.exists():
        return path.read_text(encoding="utf-8")
    raise FileNotFoundError(f"Prompt not found: {name}")


# ============================================================================
# PYDANTIC MODELS
# ============================================================================

class MatchStatus(str, Enum):
    STRONG_MATCH = "strong_match"
    GOOD_MATCH = "good_match"
    MODERATE_MATCH = "moderate_match"
    WEAK_MATCH = "weak_match"
    NO_MATCH = "no_match"


class SkillMatch(BaseModel):
    skill: str
    match_level: str  # Excellent, Strong, Good, Moderate, Weak, Missing
    evidence: str
    weight: int = Field(default=1, ge=1, le=3)  # 1=nice-to-have, 2=preferred, 3=required


class ScoreBreakdown(BaseModel):
    required_skills_score: int = Field(ge=0, le=40)
    preferred_skills_score: int = Field(ge=0, le=20)
    experience_score: int = Field(ge=0, le=20)
    domain_score: int = Field(ge=0, le=20)
    total_score: int = Field(ge=0, le=100)


class CoverLetterContent(BaseModel):
    paragraph_1: str
    paragraph_2: str
    paragraph_3: str
    paragraph_4: str


class JobFitReport(BaseModel):
    job_title: str
    company_name: Optional[str] = None

    # Deterministic scoring
    score_breakdown: ScoreBreakdown
    overall_score: int = Field(ge=0, le=100)
    match_status: MatchStatus

    recommendation: str
    summary: str

    # Skills analysis
    required_skills: List[SkillMatch]
    preferred_skills: List[SkillMatch]

    # Compact lists
    strengths: List[str]
    gaps: List[str]
    improvements: List[str]

    salary_range: Optional[str] = None
    cover_letter_content: CoverLetterContent


# ============================================================================
# DOCUMENT GENERATION
# ============================================================================

def generate_pdf(text: str) -> bytes:
    """Generate PDF from text."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    for line in text.split('\n'):
        if not line.strip():
            pdf.ln(5)
            continue

        # Header
        if line == CANDIDATE_INFO["name"]:
            pdf.set_font('Helvetica', 'B', 14)
            pdf.cell(0, 8, line, ln=True)
        # Contact lines
        elif '@' in line or 'linkedin' in line.lower():
            pdf.set_font('Helvetica', '', 9)
            pdf.cell(0, 5, line, ln=True)
        # Date
        elif '2026' in line or '2025' in line:
            pdf.ln(3)
            pdf.set_font('Helvetica', '', 11)
            pdf.cell(0, 6, line, ln=True)
            pdf.ln(3)
        # Dear/Sincerely
        elif line.startswith('Dear') or line.startswith('Sincerely'):
            pdf.set_font('Helvetica', '', 11)
            pdf.cell(0, 6, line, ln=True)
            pdf.ln(3)
        # Body
        else:
            pdf.set_font('Helvetica', '', 11)
            pdf.multi_cell(0, 6, line)
            pdf.ln(2)

    return bytes(pdf.output())


def generate_docx(text: str) -> bytes:
    """Generate DOCX from text."""
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for line in text.split('\n'):
        if not line.strip():
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        run = p.add_run(line)

        if line == CANDIDATE_INFO["name"]:
            run.bold = True
            run.font.size = Pt(14)
        elif '@' in line or 'linkedin' in line.lower():
            run.font.size = Pt(9)
        else:
            run.font.size = Pt(11)

        run.font.name = 'Times New Roman'

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_md(text: str) -> str:
    """Generate Markdown from text."""
    lines = text.split('\n')
    md = []

    for i, line in enumerate(lines):
        if not line.strip():
            md.append('')
        elif i == 0:
            md.append(f'# {line}')
        elif '@' in line or 'linkedin' in line.lower():
            md.append(f'*{line}*')
        elif line.startswith('Dear') or line.startswith('Sincerely'):
            md.append(f'**{line}**')
        else:
            md.append(line)

    md.append('\n---\n*Generated by Job Fit Analyzer*')
    return '\n'.join(md)


# ============================================================================
# COVER LETTER FORMATTING
# ============================================================================

def format_cover_letter(content: CoverLetterContent) -> str:
    """Format cover letter using external template."""
    # Prepare template data
    template_data = {
        "candidate_name": CANDIDATE_INFO["name"],
        "candidate_title": CANDIDATE_INFO["title"],
        "candidate_email": CANDIDATE_INFO["email"],
        "candidate_phone": CANDIDATE_INFO["phone"],
        "candidate_location": CANDIDATE_INFO["location"],
        "candidate_linkedin": CANDIDATE_INFO["linkedin"],
        "candidate_github": CANDIDATE_INFO["github"],
        "candidate_website": CANDIDATE_INFO["website"],
        "date": datetime.now().strftime("%B %d, %Y"),
        "paragraph_1": content.paragraph_1,
        "paragraph_2": content.paragraph_2,
        "paragraph_3": content.paragraph_3,
        "paragraph_4": content.paragraph_4
    }

    # Load and render template
    try:
        template = load_template("cover_letter.txt")
        return render_template(template, template_data)
    except FileNotFoundError:
        # Fallback to inline template
        return f"""{CANDIDATE_INFO["name"]}
{CANDIDATE_INFO["title"]}
{CANDIDATE_INFO["email"]} · {CANDIDATE_INFO["phone"]} · {CANDIDATE_INFO["location"]}
{CANDIDATE_INFO["linkedin"]} · {CANDIDATE_INFO["github"]} · {CANDIDATE_INFO["website"]}

{datetime.now().strftime("%B %d, %Y")}

{content.paragraph_1}

{content.paragraph_2}

{content.paragraph_3}

{content.paragraph_4}

Sincerely,
{CANDIDATE_INFO["name"]}
{CANDIDATE_INFO["email"]} · {CANDIDATE_INFO["linkedin"]} · {CANDIDATE_INFO["website"]}"""


# ============================================================================
# JOB PARSING
# ============================================================================

def fetch_job_description(url: str) -> str:
    """Fetch and parse job description from URL."""
    response = requests.get(url, timeout=30, headers={
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)'
    })
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)

    if len(text) < 100:
        raise ValueError("Job description too short - page may require JavaScript or login")

    return text


def read_resume(file_content: bytes, filename: str) -> str:
    """Parse resume from uploaded file."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        reader = PdfReader(BytesIO(file_content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        doc = DocxDocument(BytesIO(file_content))
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".txt":
        return file_content.decode("utf-8")
    else:
        raise ValueError(f"Unsupported format: {ext}")


# ============================================================================
# AI ANALYSIS
# ============================================================================

def get_system_prompt() -> str:
    """Load system prompt from external file."""
    try:
        return load_prompt("system_prompt.txt")
    except FileNotFoundError:
        # Fallback
        return """You are a job fit analyst. Analyze the resume against the job description.
Be evidence-based. Use specific examples from resume. Do not invent experience."""


def get_analysis_prompt(resume: str, job_desc: str) -> str:
    """Load and render analysis prompt from external file."""
    try:
        template = load_prompt("analysis_prompt.txt")
        return render_template(template, {
            "resume": resume,
            "job_description": job_desc
        })
    except FileNotFoundError:
        # Fallback
        return f"""RESUME:
{resume}

JOB DESCRIPTION:
{job_desc}

Analyze job fit with deterministic scoring. Extract job title and company name.
Identify required skills (weight=3), preferred skills (weight=2).
Calculate score breakdown using the formula.
Generate cover letter content (4 paragraphs only)."""


def create_agent() -> Agent:
    """Create analysis agent."""
    return Agent(
        "openai:gpt-4o-mini",
        output_type=JobFitReport,
        instructions=get_system_prompt()
    )


def analyze(resume: str, job_desc: str, agent: Agent) -> JobFitReport:
    """Run analysis."""
    prompt = get_analysis_prompt(resume, job_desc)
    result = agent.run_sync(prompt)
    return result.output


# ============================================================================
# STREAMLIT UI
# ============================================================================

st.set_page_config(
    page_title="Job Fit Analyzer",
    page_icon="💼",
    layout="wide"
)

# CSS
st.markdown("""
<style>
/* Cards */
.metric-card {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    border-radius: 12px;
    padding: 20px;
    text-align: center;
    color: white;
    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
}
.metric-card h1 { font-size: 3rem; margin: 0; font-weight: 700; }
.metric-card p { margin: 8px 0 0 0; opacity: 0.9; }

.status-card {
    border-radius: 12px;
    padding: 20px;
    text-align: center;
    border: 2px solid;
}
.status-strong { background: #d4edda; border-color: #28a745; }
.status-good { background: #d4edda; border-color: #5cb85c; }
.status-moderate { background: #fff3cd; border-color: #ffc107; }
.status-weak { background: #f8d7da; border-color: #dc3545; }

/* Skill Table */
.skill-table {
    width: 100%;
    border-collapse: collapse;
    font-size: 14px;
}
.skill-table th {
    background: #f8f9fa;
    padding: 10px 12px;
    text-align: left;
    border-bottom: 2px solid #dee2e6;
    font-weight: 600;
}
.skill-table td {
    padding: 8px 12px;
    border-bottom: 1px solid #eee;
    vertical-align: top;
}
.skill-table tr:hover { background: #f8f9fa; }

.badge {
    display: inline-block;
    padding: 3px 8px;
    border-radius: 12px;
    font-size: 11px;
    font-weight: 600;
}
.badge-excellent { background: #d4edda; color: #155724; }
.badge-strong { background: #d4edda; color: #155724; }
.badge-good { background: #fff3cd; color: #856404; }
.badge-moderate { background: #fff3cd; color: #856404; }
.badge-weak { background: #f8d7da; color: #721c24; }
.badge-missing { background: #f8d7da; color: #721c24; }

.badge-required { background: #dc3545; color: white; }
.badge-preferred { background: #17a2b8; color: white; }

/* Score breakdown */
.score-bar {
    background: #e9ecef;
    border-radius: 4px;
    height: 8px;
    margin: 4px 0;
}
.score-fill {
    background: linear-gradient(90deg, #667eea, #764ba2);
    height: 100%;
    border-radius: 4px;
}

/* Compact lists */
.compact-card {
    background: white;
    border-radius: 8px;
    padding: 16px;
    border-left: 4px solid;
    min-height: 180px;
}
.card-green { border-color: #28a745; background: #f8fff9; }
.card-yellow { border-color: #ffc107; background: #fffdf5; }
.card-blue { border-color: #17a2b8; background: #f5fcfd; }

/* Cover letter */
.cover-letter {
    background: white;
    border: 1px solid #dee2e6;
    border-radius: 8px;
    padding: 24px;
    font-family: 'Times New Roman', serif;
    line-height: 1.6;
    white-space: pre-line;
}
</style>
""", unsafe_allow_html=True)


def get_status_class(status: MatchStatus) -> str:
    return {
        MatchStatus.STRONG_MATCH: "status-strong",
        MatchStatus.GOOD_MATCH: "status-good",
        MatchStatus.MODERATE_MATCH: "status-moderate",
        MatchStatus.WEAK_MATCH: "status-weak",
        MatchStatus.NO_MATCH: "status-weak"
    }.get(status, "status-moderate")


def get_status_emoji(status: MatchStatus) -> str:
    return {
        MatchStatus.STRONG_MATCH: "🎯",
        MatchStatus.GOOD_MATCH: "✅",
        MatchStatus.MODERATE_MATCH: "⚠️",
        MatchStatus.WEAK_MATCH: "⚡",
        MatchStatus.NO_MATCH: "❌"
    }.get(status, "⚠️")


def render_skill_table(skills: List[SkillMatch], skill_type: str):
    """Render skills using native Streamlit components."""
    if not skills:
        st.info(f"No {skill_type} skills identified")
        return

    # Match level to emoji mapping
    level_emoji = {
        "excellent": "🟢",
        "strong": "🟢",
        "good": "🟡",
        "moderate": "🟡",
        "weak": "🔴",
        "missing": "❌"
    }

    # Header row
    col1, col2, col3 = st.columns([2, 1, 4])
    with col1:
        st.markdown("**Skill**")
    with col2:
        st.markdown("**Match**")
    with col3:
        st.markdown("**Evidence**")

    st.divider()

    # Data rows
    for s in skills:
        col1, col2, col3 = st.columns([2, 1, 4])
        emoji = level_emoji.get(s.match_level.lower(), "⚪")

        with col1:
            st.markdown(f"**{s.skill}**")
        with col2:
            st.markdown(f"{emoji} {s.match_level}")
        with col3:
            evidence = s.evidence[:120] + "..." if len(s.evidence) > 120 else s.evidence
            st.caption(evidence)


class FullReport(BaseModel):
    """Complete report with all input and output data."""
    generated_at: str
    version: str = "1.0"

    # Input
    input_job_url: str
    input_job_description_raw: str
    input_resume_file: str
    input_resume_text: str

    # Analysis
    analysis: Dict[str, Any]

    # Cover letter
    cover_letter_formatted: str

    # Candidate info
    candidate: Dict[str, str]


def render_dashboard(report: JobFitReport, cover_letter: str, job_url: str = "", job_desc: str = "", resume_text: str = "", resume_filename: str = ""):
    """Render the analysis dashboard."""

    # Header
    st.markdown(f"## 💼 {report.job_title}")
    if report.company_name:
        st.caption(f"at {report.company_name}")
    st.markdown("---")

    # KPI Row
    col1, col2, col3 = st.columns([1, 2, 1])

    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <h1>{report.overall_score}</h1>
            <p>Match Score</p>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        status_class = get_status_class(report.match_status)
        status_emoji = get_status_emoji(report.match_status)
        status_label = report.match_status.value.replace("_", " ").title()

        st.markdown(f"""
        <div class="status-card {status_class}">
            <h2 style="margin: 0;">{status_emoji} {status_label}</h2>
            <p style="margin: 10px 0 0 0; color: #333;">{report.recommendation}</p>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        if report.salary_range:
            st.markdown(f"""
            <div class="metric-card" style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);">
                <p style="margin: 0; font-size: 14px;">💰 Salary</p>
                <h3 style="margin: 8px 0 0 0; font-size: 1.3rem;">{report.salary_range}</h3>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style="background: #f8f9fa; border-radius: 12px; padding: 20px; text-align: center; color: #6c757d;">
                <p style="margin: 0;">💰 Salary</p>
                <p style="margin: 8px 0 0 0;">Not specified</p>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Score Breakdown
    st.markdown("### 📊 Score Breakdown")

    breakdown = report.score_breakdown
    cols = st.columns(4)

    scores = [
        ("Required Skills", breakdown.required_skills_score, 40),
        ("Preferred Skills", breakdown.preferred_skills_score, 20),
        ("Experience", breakdown.experience_score, 20),
        ("Domain Match", breakdown.domain_score, 20)
    ]

    for i, (label, score, max_score) in enumerate(scores):
        with cols[i]:
            pct = (score / max_score) * 100
            st.markdown(f"""
            <div style="background: white; padding: 12px; border-radius: 8px; border: 1px solid #eee;">
                <div style="display: flex; justify-content: space-between; margin-bottom: 4px;">
                    <span style="font-size: 13px; color: #666;">{label}</span>
                    <span style="font-weight: 600;">{score}/{max_score}</span>
                </div>
                <div class="score-bar"><div class="score-fill" style="width: {pct}%;"></div></div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Summary
    st.info(f"**Summary:** {report.summary}")

    # Quick Overview - 3 columns
    st.markdown("### 📋 Quick Overview")
    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("**💪 Strengths**")
        items = "".join([f"<div style='margin: 6px 0;'>✓ {s}</div>" for s in report.strengths])
        st.markdown(f'<div class="compact-card card-green">{items}</div>', unsafe_allow_html=True)

    with col2:
        st.markdown("**⚠️ Gaps**")
        items = "".join([f"<div style='margin: 6px 0;'>• {g}</div>" for g in report.gaps]) or "<div>None identified</div>"
        st.markdown(f'<div class="compact-card card-yellow">{items}</div>', unsafe_allow_html=True)

    with col3:
        st.markdown("**💡 Improvements**")
        items = "".join([f"<div style='margin: 6px 0;'>→ {i}</div>" for i in report.improvements])
        st.markdown(f'<div class="compact-card card-blue">{items}</div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Skills Tables
    st.markdown("### 🎯 Skills Analysis")

    tab1, tab2 = st.tabs(["Required Skills", "Preferred Skills"])

    with tab1:
        render_skill_table(report.required_skills, "required")

    with tab2:
        render_skill_table(report.preferred_skills, "preferred")

    st.markdown("---")

    # Cover Letter
    st.markdown("### 📝 Cover Letter")

    col1, col2 = st.columns([3, 1])

    with col1:
        st.markdown(f'<div class="cover-letter">{cover_letter}</div>', unsafe_allow_html=True)

    with col2:
        st.markdown("**📥 Downloads**")

        ts = datetime.now().strftime('%Y%m%d_%H%M%S')

        st.download_button("📝 Markdown", generate_md(cover_letter),
                          f"cover_letter_{ts}.md", "text/markdown", use_container_width=True)

        st.download_button("📕 PDF", generate_pdf(cover_letter),
                          f"cover_letter_{ts}.pdf", "application/pdf", use_container_width=True)

        st.download_button("📘 Word", generate_docx(cover_letter),
                          f"cover_letter_{ts}.docx",
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          use_container_width=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # Full report with all data
        full_report = FullReport(
            generated_at=datetime.now().isoformat(),
            input_job_url=job_url,
            input_job_description_raw=job_desc,
            input_resume_file=resume_filename,
            input_resume_text=resume_text,
            analysis=report.model_dump(),
            cover_letter_formatted=cover_letter,
            candidate=CANDIDATE_INFO
        )
        report_json = json.dumps(full_report.model_dump(), indent=2, default=str)
        st.download_button("📊 JSON Report", report_json,
                          f"report_{ts}.json", "application/json", use_container_width=True)


def main():
    """Main app."""

    # Session state
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "resume_text" not in st.session_state:
        st.session_state.resume_text = None
    if "resume_filename" not in st.session_state:
        st.session_state.resume_filename = None
    if "agent" not in st.session_state:
        st.session_state.agent = create_agent()

    # Sidebar
    with st.sidebar:
        st.title("📄 Resume")

        uploaded = st.file_uploader("Upload resume", type=["pdf", "docx", "txt"])

        if uploaded:
            try:
                content = uploaded.read()
                st.session_state.resume_text = read_resume(content, uploaded.name)
                st.session_state.resume_filename = uploaded.name
                st.success(f"✓ {uploaded.name}")
            except Exception as e:
                st.error(f"Error: {e}")

        st.divider()

        if st.button("🗑️ Clear", use_container_width=True):
            st.session_state.messages = []
            st.rerun()

        st.divider()
        st.caption("**Scoring Formula:**")
        st.caption("• Required Skills: 40%")
        st.caption("• Preferred Skills: 20%")
        st.caption("• Experience: 20%")
        st.caption("• Domain Match: 20%")

    # Main
    st.title("💼 Job Fit Analyzer")

    if not st.session_state.resume_text:
        st.info("👋 Upload your resume to get started")
        return

    # Chat history
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            if msg["type"] == "url":
                st.markdown(f"Analyze: [{msg['content']}]({msg['content']})")
            elif msg["type"] == "report":
                render_dashboard(
                    JobFitReport(**msg["content"]["report"]),
                    msg["content"]["cover_letter"],
                    msg["content"].get("job_url", ""),
                    msg["content"].get("job_desc", ""),
                    msg["content"].get("resume_text", ""),
                    msg["content"].get("resume_filename", "")
                )

    # Input
    if url := st.chat_input("Paste job URL..."):
        # Validate
        try:
            result = urlparse(url)
            if not all([result.scheme, result.netloc]):
                st.error("Invalid URL format")
                return
        except:
            st.error("Invalid URL")
            return

        # Add message
        st.session_state.messages.append({"role": "user", "content": url, "type": "url"})

        with st.chat_message("user"):
            st.markdown(f"Analyze: [{url}]({url})")

        with st.chat_message("assistant"):
            try:
                with st.status("Analyzing...", expanded=True) as status:
                    status.update(label="📥 Fetching job posting...")
                    try:
                        job_desc = fetch_job_description(url)
                    except Exception as e:
                        status.update(label="❌ Failed", state="error")
                        st.error(f"""**Failed to fetch job posting**

Error: {str(e)}

**Possible causes:**
- Page requires login
- JavaScript-rendered content
- Bot protection / CAPTCHA
- Invalid URL

**Solution:** Copy the job description text manually.""")
                        return

                    time.sleep(0.5)
                    status.update(label="📊 Analyzing skills...")
                    time.sleep(1.0)
                    status.update(label="🤖 Running AI analysis...")

                    report = analyze(
                        st.session_state.resume_text,
                        job_desc,
                        st.session_state.agent
                    )

                    time.sleep(0.5)
                    status.update(label="📝 Generating cover letter...")
                    cover_letter = format_cover_letter(report.cover_letter_content)

                    status.update(label="✅ Complete!", state="complete")

                render_dashboard(
                    report, cover_letter, url, job_desc,
                    st.session_state.resume_text,
                    st.session_state.resume_filename or ""
                )

                st.session_state.messages.append({
                    "role": "assistant",
                    "content": {
                        "report": report.model_dump(),
                        "cover_letter": cover_letter,
                        "job_url": url,
                        "job_desc": job_desc,
                        "resume_text": st.session_state.resume_text,
                        "resume_filename": st.session_state.resume_filename or ""
                    },
                    "type": "report"
                })

            except Exception as e:
                st.error(f"Analysis error: {e}")
                import traceback
                with st.expander("Details"):
                    st.code(traceback.format_exc())


if __name__ == "__main__":
    # Load .env file
    from dotenv import load_dotenv
    load_dotenv()

    # Support both OPENAI_API_KEY and API_KEY
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("API_KEY")
    if api_key:
        os.environ["OPENAI_API_KEY"] = api_key

    if not os.getenv("OPENAI_API_KEY"):
        st.error("Set OPENAI_API_KEY or API_KEY in .env file")
    else:
        main()
